import os
import re
import shutil
import tempfile
import logging
from docx import Document
from docxcompose.composer import Composer
from natsort import natsorted
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 获取 main.py 中定义的 logger
logger = logging.getLogger("WordMergerLogger")


def process_document_headings(doc, display_name):
    """
    核心逻辑：将文档所有标题降级，并在开头插入自定义名称作为一级标题，同时切断自动编号连号
    """
    first_num_id = None

    for p in doc.paragraphs:
        style_name = p.style.name
        match_en = re.match(r'^Heading (\d)$', style_name, re.IGNORECASE)
        match_zh = re.match(r'^标题\s*(\d)$', style_name)

        level = None
        if match_en:
            level = int(match_en.group(1))
        elif match_zh:
            level = int(match_zh.group(1))

        if level and level < 9:
            # --- 【步骤 A】 提取该文档原始的列表编号 ID ---
            if first_num_id is None:
                pPr = p._element.pPr
                numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
                if numPr is None and p.style is not None and p.style._element.pPr is not None:
                    numPr = p.style._element.pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_elem = numPr.find(qn('w:numId'))
                    if numId_elem is not None:
                        first_num_id = numId_elem.get(qn('w:val'))

            # --- 【步骤 B】 样式外观降级 ---
            target_en = f'Heading {level + 1}'
            target_zh = f'标题 {level + 1}'
            try:
                p.style = target_en
            except KeyError:
                try:
                    p.style = target_zh
                except KeyError:
                    pass

            # --- 【步骤 C】 底层列表层级 (ilvl) 强制降级 ---
            pPr = p._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                ilvl_elem = numPr.find(qn('w:ilvl'))
                if ilvl_elem is not None:
                    current_ilvl = int(ilvl_elem.get(qn('w:val'), '0'))
                    ilvl_elem.set(qn('w:val'), str(current_ilvl + 1))
                else:
                    new_ilvl = OxmlElement('w:ilvl')
                    new_ilvl.set(qn('w:val'), '1')
                    numPr.append(new_ilvl)

    # --- 【步骤 D】 插入最高级标题（作为断层切割器） ---
    heading_style = 'Heading 1'
    if '标题 1' in doc.styles and 'Heading 1' not in doc.styles:
        heading_style = '标题 1'

    if len(doc.paragraphs) > 0:
        try:
            new_p = doc.paragraphs[0].insert_paragraph_before(display_name, style=heading_style)
        except KeyError:
            new_p = doc.paragraphs[0].insert_paragraph_before(display_name)
    else:
        new_p = doc.add_heading(display_name, level=1)

    # --- 【步骤 E】 把抓取到的编号 ID 赋予最高级标题，彻底打断连号！ ---
    if first_num_id is not None:
        pPr = new_p._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            numPr = OxmlElement('w:numPr')
            pPr.append(numPr)
        else:
            for child in list(numPr):
                if child.tag in [qn('w:ilvl'), qn('w:numId')]:
                    numPr.remove(child)

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), first_num_id)

        numPr.append(ilvl)
        numPr.append(numId)

    return doc


def core_merge_logic(selected_items, save_path, progress_callback):
    """
    分治合并逻辑引擎
    :param selected_items: 文件对象列表 [{'path':..., 'displayName':...}]
    :param save_path: 保存路径
    :param progress_callback: 进度回调函数，用于更新前端 UI
    """
    total_files = len(selected_items)
    logger.info(f"=== 开始执行合并任务，计划合并: {total_files} 个文件 ===")

    if total_files == 0:
        return {"status": "error", "msg": "没有需要合并的文件！"}

    if total_files == 1:
        shutil.copy2(selected_items[0]['path'], save_path)
        progress_callback(100, "合并完成！")
        return {"status": "success", "msg": "只有一个文件，已直接复制到目标位置！"}

    BATCH_SIZE = 15
    temp_dir = tempfile.mkdtemp()
    temp_file_paths = []

    try:
        # 1. 分批处理
        batches = [selected_items[i:i + BATCH_SIZE] for i in range(0, total_files, BATCH_SIZE)]
        processed_count = 0

        for batch_idx, batch in enumerate(batches):
            temp_file_path = os.path.join(temp_dir, f"batch_{batch_idx}.docx")

            # 取每一组的第一个作为 Master
            master_item = batch[0]
            master_doc = Document(master_item['path'])
            master_doc = process_document_headings(master_doc, master_item['displayName'])
            composer = Composer(master_doc)
            processed_count += 1

            for item in batch[1:]:
                percent = int((processed_count / total_files) * 80)
                progress_callback(percent,
                                  f"[分批加速] 正在处理编号 ({processed_count}/{total_files}): {item['displayName']}")

                doc_to_append = Document(item['path'])
                doc_to_append = process_document_headings(doc_to_append, item['displayName'])
                composer.append(doc_to_append)
                processed_count += 1

            composer.save(temp_file_path)
            temp_file_paths.append(temp_file_path)

        # 2. 组装分块
        progress_callback(80, "各子文档目录处理完毕，开始组装最终文档...")
        final_master_doc = Document(temp_file_paths[0])
        final_composer = Composer(final_master_doc)

        for i, t_file in enumerate(temp_file_paths[1:]):
            percent = 80 + int(((i + 1) / len(temp_file_paths)) * 15)
            progress_callback(percent, f"正在组装最终分块 ({i + 2}/{len(temp_file_paths)})...")
            doc_to_append = Document(t_file)
            final_composer.append(doc_to_append)

        # 3. 最终写入
        progress_callback(95, "正在写入最终文件...")
        final_composer.save(save_path)
        progress_callback(100, "合并完成！")

        return {"status": "success", "msg": f"🎉 成功合并了 {total_files} 个文件！"}

    except Exception as e:
        logger.error(f"合并崩溃: {str(e)}", exc_info=True)
        return {"status": "error", "msg": f"发生错误: {str(e)}"}
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)