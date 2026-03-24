import os
import re
import eel
import tkinter as tk
import tempfile
import shutil
import logging
from logging.handlers import TimedRotatingFileHandler
from tkinter import filedialog
from docx import Document
from docxcompose.composer import Composer
from natsort import natsorted
from eel import sleep as eel_sleep

# ===== 核心新增：引入处理 Word 底层 XML 的库 =====
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===============================================

# ================= 1. 日志系统配置 =================
LOG_FILE = "word_merger_log.txt"
logger = logging.getLogger("WordMergerLogger")
logger.setLevel(logging.INFO)

file_handler = TimedRotatingFileHandler(filename=LOG_FILE, when="D", interval=1, backupCount=3, encoding="utf-8")
formatter = logging.Formatter('[%(asctime)s] [%(levelname)s] - %(message)s', datefmt='%Y-%m-%d %H:%M:%S')
file_handler.setFormatter(formatter)
logger.addHandler(file_handler)

console_handler = logging.StreamHandler()
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

logger.info("========================================")
logger.info("程序启动：Word 极速合并助手 (彻底修复连号版) 已打开")

# 初始化 Eel
eel.init('dist')


# ================= 2. 交互与扫描逻辑 =================
@eel.expose
def py_choose_and_scan():
    logger.info("用户点击了【选择文件夹】按钮。")

    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    root.lift()

    folder_path = filedialog.askdirectory(parent=root, title="选择存放 Word 文件的文件夹")
    root.destroy()

    if not folder_path:
        logger.info("用户取消了文件夹选择。")
        return None

    eel.set_scanning_status(True)()
    logger.info(f"开始深度扫描文件夹: {folder_path}...")

    temp_files = []
    file_count = 0
    for root_dir, _, files in os.walk(folder_path):
        for file in files:
            file_count += 1
            if file_count % 500 == 0:
                eel_sleep(0.01)

            if file.endswith('.docx') and not file.startswith('~$'):
                temp_files.append(os.path.join(root_dir, file))

    sorted_files = natsorted(temp_files)
    logger.info(f"扫描完毕，共找到 {len(sorted_files)} 个有效文件。")

    return {"folder_path": folder_path, "files": sorted_files}


@eel.expose
def py_choose_save_path():
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    root.lift()

    save_path = filedialog.asksaveasfilename(
        parent=root,
        title="保存合并后的文件",
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        initialfile="合并完成的文档.docx"
    )
    root.destroy()
    if save_path:
        logger.info(f"用户选择了保存路径: {save_path}")
    return save_path


# ================= 3. 核心：目录降级与防连号处理 =================
def process_document_headings(doc, display_name):
    """将文档所有标题降一级，并在开头插入自定义名称作为一级标题，同时切断自动编号连号"""
    first_num_id = None

    for p in doc.paragraphs:
        style_name = p.style.name
        match_en = re.match(r'^Heading (\d)$', style_name, re.IGNORECASE)
        match_zh = re.match(r'^标题\s*(\d)$', style_name)

        level = None
        if match_en:
            level = int(match_en.group(1))
        elif match_zh:
            level = int(match_zh.group(1))

        if level and level < 9:
            # --- 【步骤 A】 提取该文档原始的列表编号 ID ---
            if first_num_id is None:
                pPr = p._element.pPr
                numPr = pPr.find(qn('w:numPr')) if pPr is not None else None
                # 如果段落没有直接的编号，去它的样式里找
                if numPr is None and p.style is not None and p.style._element.pPr is not None:
                    numPr = p.style._element.pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_elem = numPr.find(qn('w:numId'))
                    if numId_elem is not None:
                        first_num_id = numId_elem.get(qn('w:val'))

            # --- 【步骤 B】 样式外观降级 ---
            target_en = f'Heading {level + 1}'
            target_zh = f'标题 {level + 1}'
            try:
                p.style = target_en
            except KeyError:
                try:
                    p.style = target_zh
                except KeyError:
                    pass

            # --- 【步骤 C】 底层列表层级 (ilvl) 强制降级 ---
            pPr = p._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                ilvl_elem = numPr.find(qn('w:ilvl'))
                if ilvl_elem is not None:
                    current_ilvl = int(ilvl_elem.get(qn('w:val'), '0'))
                    ilvl_elem.set(qn('w:val'), str(current_ilvl + 1))
                else:
                    new_ilvl = OxmlElement('w:ilvl')
                    new_ilvl.set(qn('w:val'), '1')
                    numPr.append(new_ilvl)

    # --- 【步骤 D】 插入最高级标题（作为断层切割器） ---
    heading_style = 'Heading 1'
    if '标题 1' in doc.styles and 'Heading 1' not in doc.styles:
        heading_style = '标题 1'

    if len(doc.paragraphs) > 0:
        try:
            new_p = doc.paragraphs[0].insert_paragraph_before(display_name, style=heading_style)
        except KeyError:
            new_p = doc.paragraphs[0].insert_paragraph_before(display_name)
    else:
        new_p = doc.add_heading(display_name, level=1)

    # --- 【步骤 E】 把抓取到的编号 ID 赋予最高级标题，彻底打断连号！ ---
    if first_num_id is not None:
        pPr = new_p._element.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            numPr = OxmlElement('w:numPr')
            pPr.append(numPr)
        else:
            # 清理旧的层级设置
            for child in list(numPr):
                if child.tag in [qn('w:ilvl'), qn('w:numId')]:
                    numPr.remove(child)

        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')  # 设为多级列表的根层级
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), first_num_id)

        numPr.append(ilvl)
        numPr.append(numId)

    return doc


# ================= 4. 核心：合并逻辑引擎 =================
@eel.expose
def py_merge_files(selected_items, save_path):
    total_files = len(selected_items)
    logger.info(f"=== 开始执行合并任务，计划合并: {total_files} 个文件 ===")

    if total_files == 0:
        return {"status": "error", "msg": "没有需要合并的文件！"}

    if total_files == 1:
        shutil.copy2(selected_items[0]['path'], save_path)
        eel.update_progress(100, "合并完成！")()
        return {"status": "success", "msg": "只有一个文件，已直接复制到目标位置！"}

    BATCH_SIZE = 15

    if total_files <= BATCH_SIZE:
        logger.info("文件数量较少，进入【常规合并模式】...")
        return merge_standard(selected_items, save_path)

    logger.info("文件数量较多，自动触发【分治合并加速模式】...")
    temp_dir = tempfile.mkdtemp()
    temp_file_paths = []

    try:
        batches = [selected_items[i:i + BATCH_SIZE] for i in range(0, total_files, BATCH_SIZE)]
        processed_count = 0

        for batch_idx, batch in enumerate(batches):
            temp_file_path = os.path.join(temp_dir, f"batch_{batch_idx}.docx")

            master_item = batch[0]
            master_doc = Document(master_item['path'])
            master_doc = process_document_headings(master_doc, master_item['displayName'])
            composer = Composer(master_doc)
            processed_count += 1

            for doc_idx, item in enumerate(batch[1:]):
                percent = int((processed_count / total_files) * 80)
                eel.update_progress(percent,
                                    f"[分批加速] 正在处理与重置目录编号 ({processed_count}/{total_files}): {item['displayName']}")()

                doc_to_append = Document(item['path'])
                doc_to_append = process_document_headings(doc_to_append, item['displayName'])
                composer.append(doc_to_append)
                processed_count += 1

            composer.save(temp_file_path)
            temp_file_paths.append(temp_file_path)

        eel.update_progress(80, "各子文档目录处理完毕，开始组装最终文档...")()
        logger.info("预合并完成，开始组装最终大文件...")

        final_master_doc = Document(temp_file_paths[0])
        final_composer = Composer(final_master_doc)

        for i, t_file in enumerate(temp_file_paths[1:]):
            percent = 80 + int(((i + 1) / len(temp_file_paths)) * 15)
            eel.update_progress(percent, f"正在组装最终分块 ({i + 2}/{len(temp_file_paths)})...")()
            doc_to_append = Document(t_file)
            final_composer.append(doc_to_append)

        eel.update_progress(95, "正在写入最终文件...")()
        final_composer.save(save_path)

        eel.update_progress(100, "合并完成！")()
        logger.info("=== 合并与目录重组任务圆满成功 ===")
        return {"status": "success",
                "msg": f"🎉 成功合并了 {total_files} 个文件！\n\n目录重组完成，编号已成功切断重置。\n已保存至:\n{save_path}"}

    except Exception as e:
        logger.error(f"合并崩溃: {str(e)}", exc_info=True)
        return {"status": "error", "msg": f"发生错误: {str(e)}"}
    finally:
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir, ignore_errors=True)
            logger.info("已清理临时工作目录。")


def merge_standard(selected_items, save_path):
    try:
        total_files = len(selected_items)
        eel.update_progress(0, f"正在初始化主文档目录...")()

        master_item = selected_items[0]
        master_doc = Document(master_item['path'])
        master_doc = process_document_headings(master_doc, master_item['displayName'])
        composer = Composer(master_doc)

        for i, item in enumerate(selected_items[1:]):
            percent = ((i + 1) / total_files) * 100
            eel.update_progress(percent, f"正在处理目录编号 ({i + 2}/{total_files}): {item['displayName']}")()

            doc_to_append = Document(item['path'])
            doc_to_append = process_document_headings(doc_to_append, item['displayName'])
            composer.append(doc_to_append)

        eel.update_progress(95, "正在保存合并文档...")()
        composer.save(save_path)
        eel.update_progress(100, "合并完成！")()
        return {"status": "success",
                "msg": f"🎉 成功合并了 {total_files} 个文件！\n\n目录已自动重组且编号重启。\n已保存至:\n{save_path}"}
    except Exception as e:
        logger.error(f"常规合并发生错误: {str(e)}", exc_info=True)
        return {"status": "error", "msg": str(e)}


if __name__ == '__main__':
    try:
        # 修改了打开窗口的默认比例，适配长列表浏览
        eel.start('index.html', size=(800, 850), port=0)
    except (SystemExit, MemoryError, KeyboardInterrupt):
        logger.info("程序正常关闭。")
    except Exception as e:
        logger.critical(f"程序遭遇灾难性错误而退出: {str(e)}", exc_info=True)